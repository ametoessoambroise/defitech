"""
CareerCraft AI - Module de génération et d'optimisation de CV
Gère la création, l'édition et l'export de CV avec l'IA
"""

import os
import json
import logging
import re
from datetime import datetime
from typing import Dict, Optional, Any, List
from pathlib import Path

from flask import Blueprint, request, jsonify, send_file, current_app
from flask_login import current_user, login_required
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.gemini_integration import GeminiIntegration

from app.models.user import User
from app.models.etudiant import Etudiant
from app.models.note import Note
from app.models.matiere import Matiere
from app.models.competence import Competence
from app.models.experience import Experience
from app.models.formation import Formation
from app.models.langue import Langue
from app.models.projet import Projet

# Configuration du logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Création du blueprint
career_craft_bp = Blueprint("career_craft", __name__, url_prefix="/api/career-craft")


@career_craft_bp.route("/career_craft", methods=["GET"])
def career_craft():
    """Route principale pour CareerCraft - redirection vers l'interface"""
    return jsonify(
        {
            "success": True,
            "message": "CareerCraft API - Endpoint principal",
            "endpoints": {
                "templates": "/api/career-craft/templates",
                "generate": "/api/career-craft/generate",
                "download": "/api/career-craft/download",
                "preview": "/api/career-craft/preview",
                "data": "/api/career-craft/data",
                "optimize": "/api/career-craft/optimize",
                "export_history": "/api/career-craft/export/history",
                "delete": "/api/career-craft/delete/<filename>",
                "suggestions": "/api/career-craft/suggestions",
                "health": "/api/career-craft/health",
            },
        }
    )


# Configuration des dossiers
BASE_DIR = Path("static")
UPLOAD_FOLDER = BASE_DIR / "uploads" / "cv"
TEMPLATES_FOLDER = BASE_DIR / "templates" / "cv"
OUTPUT_FOLDER = BASE_DIR / "exports" / "cv"

# Création des dossiers s'ils n'existent pas
for folder in [UPLOAD_FOLDER, TEMPLATES_FOLDER, OUTPUT_FOLDER]:
    folder.mkdir(parents=True, exist_ok=True)


class CVGenerator:
    """Classe principale pour la génération et la gestion des CV"""

    # Constantes
    SKILL_LEVELS = {
        (16, float("inf")): "Expert",
        (14, 16): "Avancé",
        (12, 14): "Intermédiaire",
        (10, 12): "Débutant avancé",
        (0, 10): "Débutant",
    }

    def __init__(self, user_id: int):
        self.user_id = user_id
        self.user = self._load_user(user_id)
        self.etudiant = self._load_etudiant(user_id)
        self.templates = self._load_templates()
        self.gemini = self._init_gemini()
        self.styles = self._init_styles()

    def _load_user(self, user_id: int) -> Optional[User]:
        """Charge l'utilisateur avec gestion d'erreur"""
        try:
            user = User.query.get(user_id)
            if not user:
                logger.error(f"Utilisateur non trouvé: ID {user_id}")
            else:
                logger.info(
                    f"Utilisateur chargé: {user.prenom} {user.nom} (ID: {user_id})"
                )
            return user
        except Exception as e:
            logger.error(f"Erreur lors du chargement de l'utilisateur {user_id}: {e}")
            return None

    def _load_etudiant(self, user_id: int) -> Optional[Etudiant]:
        """Charge le profil étudiant avec gestion d'erreur"""
        try:
            etudiant = Etudiant.query.filter_by(user_id=user_id).first()
            if not etudiant:
                logger.warning(
                    f"Profil étudiant non trouvé pour l'utilisateur {user_id}"
                )
            else:
                logger.info(f"Profil étudiant chargé: ID {etudiant.id}")
            return etudiant
        except Exception as e:
            logger.error(f"Erreur lors du chargement du profil étudiant: {e}")
            return None

    def _init_gemini(self) -> Optional[GeminiIntegration]:
        """Initialise l'intégration Gemini"""
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                logger.warning("Clé API Gemini non configurée")
                return None
            return GeminiIntegration(api_key=api_key)
        except Exception as e:
            logger.error(f"Erreur lors de l'initialisation de Gemini: {e}")
            return None

    def _init_styles(self) -> Dict[str, Dict]:
        """Initialise les styles de document"""
        return {
            "title": {"font": "Helvetica-Bold", "size": 16, "color": (0.2, 0.2, 0.6)},
            "header": {"font": "Helvetica-Bold", "size": 14, "color": (0.2, 0.2, 0.2)},
            "subheader": {
                "font": "Helvetica-Bold",
                "size": 12,
                "color": (0.3, 0.3, 0.3),
            },
            "normal": {"font": "Helvetica", "size": 11, "color": (0.1, 0.1, 0.1)},
            "small": {"font": "Helvetica", "size": 10, "color": (0.4, 0.4, 0.4)},
        }

    def _load_templates(self) -> Dict[str, Dict]:
        """Charge les modèles de CV depuis le dossier des templates"""
        templates = {}
        try:
            for file_path in TEMPLATES_FOLDER.glob("*.json"):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        template_data = json.load(f)
                        template_id = template_data.get("id")
                        if template_id:
                            templates[template_id] = template_data
                            logger.info(f"Template chargé: {template_id}")
                except json.JSONDecodeError as e:
                    logger.error(f"Erreur JSON dans {file_path.name}: {e}")
                except Exception as e:
                    logger.error(f"Erreur lors du chargement de {file_path.name}: {e}")
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du dossier templates: {e}")

        return templates

    def get_user_data(self) -> Dict[str, Any]:
        """Récupère toutes les données de l'utilisateur nécessaires à la génération du CV"""
        logger.info(f"Récupération des données pour l'utilisateur {self.user_id}")

        # Validation de base
        if not self.user:
            return {"error": "Utilisateur non trouvé"}

        if not self.etudiant:
            return {
                "error": "Profil étudiant non trouvé",
                "user": {
                    "id": self.user.id,
                    "nom": self.user.nom,
                    "prenom": self.user.prenom,
                    "email": self.user.email,
                    "role": self.user.role,
                },
            }

        try:
            # Récupération des données
            competences = self._get_competences()
            competences_techniques = self._get_competences_techniques()
            experiences = self._get_experiences()
            formations = self._get_formations()
            langues = self._get_langues()
            projets = self._get_projets()
            certifications = self._get_certifications()
            references = self._get_references()
            centres_interet = self._get_centres_interet()

            data = {
                "utilisateur": self._format_user_info(),
                "competences": competences,
                "competences_techniques": competences_techniques,
                "experiences": experiences,
                "formations": formations,
                "langues": langues,
                "projets": projets,
                "certifications": certifications,
                "references": references,
                "centres_interet": centres_interet,
            }

            # Ajouter la complétion et les suggestions
            data["completion"] = self.calculate_profile_completion()
            data["suggestions"] = self.get_suggestions(data)

            return data
        except Exception as e:
            logger.error(
                f"Erreur lors de la récupération des données: {e}", exc_info=True
            )
            return {"error": f"Erreur lors de la récupération des données: {str(e)}"}

    def _format_user_info(self) -> Dict[str, str]:
        """Formate les informations de l'utilisateur"""
        return {
            "id": self.user.id,
            "nom": self.user.nom or "",
            "prenom": self.user.prenom or "",
            "email": self.user.email or "",
            "telephone": self.user.telephone or "",
            "adresse": self.user.adresse or "",
            "ville": self.user.ville or "",
            "code_postal": self.user.code_postal or "",
            "pays": self.user.pays or "",
            "date_naissance": (
                self.user.date_naissance.strftime("%d/%m/%Y")
                if hasattr(self.user, "date_naissance") and self.user.date_naissance
                else ""
            ),
            "linkedin": getattr(self.user, "linkedin", "") or "",
            "github": getattr(self.user, "github", "") or "",
            "bio": getattr(self.user, "bio", "") or "",
            "photo_profil": getattr(self.user, "photo_profil", "") or "",
            "titre_professionnel": getattr(self.user, "titre_professionnel", "") or "",
            "disponibilite": getattr(self.user, "disponibilite", "") or "",
            "permis_conduire": getattr(self.user, "permis_conduire", "") or "",
            "pretentions_salariales": getattr(self.user, "pretentions_salariales", "") or "",
        }

    def calculate_profile_completion(self) -> Dict[str, Any]:
        """Calcule le pourcentage de complétion du profil"""
        score = 0
        details = {}

        # 1. Infos de base (30%)
        basic_fields = ["nom", "prenom", "email", "telephone", "adresse", "bio"]
        basic_score = 0
        for field in basic_fields:
            if getattr(self.user, field, None):
                basic_score += 5
        score += basic_score
        details["infos_base"] = basic_score

        # 2. Expériences (20%)
        exp_count = Experience.query.filter_by(user_id=self.user_id).count()
        exp_score = min(exp_count * 10, 20)
        score += exp_score
        details["experiences"] = exp_score

        # 3. Formations (20%)
        form_count = Formation.query.filter_by(user_id=self.user_id).count()
        form_score = min(form_count * 10, 20)
        score += form_score
        details["formations"] = form_score

        # 4. Compétences (15%)
        comp_count = Competence.query.filter_by(user_id=self.user_id).count()
        comp_score = min(comp_count * 5, 15)
        score += comp_score
        details["competences"] = comp_score

        # 5. Langues (5%)
        lang_count = Langue.query.filter_by(user_id=self.user_id).count()
        lang_score = 5 if lang_count > 0 else 0
        score += lang_score
        details["langues"] = lang_score

        # 6. Projets (10%)
        proj_count = Projet.query.filter_by(user_id=self.user_id).count()
        proj_score = min(proj_count * 5, 10)
        score += proj_score
        details["projets"] = proj_score

        return {"score": score, "details": details}

    def get_suggestions(self, data: Dict) -> List[str]:
        """Génère des suggestions pour améliorer le profil"""
        suggestions = []
        user = data["utilisateur"]

        if not user.get("bio"):
            suggestions.append(
                "Ajoutez une biographie pour vous présenter aux recruteurs."
            )
        if not user.get("telephone"):
            suggestions.append(
                "Ajoutez votre numéro de téléphone pour être recontacté."
            )
        if not data.get("experiences"):
            suggestions.append("Ajoutez au moins une expérience professionnelle.")
        if not data.get("formations"):
            suggestions.append("Précisez votre parcours académique.")
        if len(data.get("competences", [])) < 3:
            suggestions.append(
                "Ajoutez plus de compétences pour valoriser votre profil."
            )
        if not data.get("projets"):
            suggestions.append("Mettez en avant vos projets personnels ou académiques.")

        return suggestions

    def _get_competences(self) -> List[Dict]:
        """Récupère les compétences générales"""
        try:
            competences = Competence.query.filter_by(
                user_id=self.etudiant.user_id
            ).all()
            return [
                {
                    "id": c.id,
                    "nom": c.nom,
                    "niveau": c.niveau,
                    "categorie": getattr(c, "categorie", "Général"),
                }
                for c in competences
            ]
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des compétences: {e}")
            return []

    def _get_competences_techniques(self) -> List[Dict]:
        """Calcule les compétences techniques à partir des notes"""
        try:
            matieres = (
                Matiere.query.join(Note)
                .filter(Note.etudiant_id == self.etudiant.id)
                .distinct()
                .all()
            )

            competences_techniques = []
            for matiere in matieres:
                notes = Note.query.filter_by(
                    etudiant_id=self.etudiant.id, matiere_id=matiere.id
                ).all()

                if notes:
                    moyenne = sum(note.note for note in notes) / len(notes)
                    niveau = self._get_skill_level(moyenne)
                    categorie = (
                        matiere.filiere.nom
                        if hasattr(matiere, "filiere") and matiere.filiere
                        else "Autre"
                    )

                    competences_techniques.append(
                        {"nom": matiere.nom, "niveau": niveau, "categorie": categorie}
                    )

            return competences_techniques
        except Exception as e:
            logger.error(f"Erreur lors du calcul des compétences techniques: {e}")
            return []

    def _get_experiences(self) -> List[Dict]:
        """Récupère les expériences professionnelles"""
        try:
            experiences = (
                Experience.query.filter_by(user_id=self.etudiant.user_id)
                .order_by(Experience.date_debut.desc())
                .all()
            )

            return [
                {
                    "id": e.id,
                    "titre": e.poste or "",  # Corrigé: e.titre -> e.poste
                    "entreprise": e.entreprise or "",
                    "lieu": e.lieu or "",
                    "date_debut": (
                        e.date_debut.strftime("%m/%Y") if e.date_debut else ""
                    ),
                    "date_fin": (
                        e.date_fin.strftime("%m/%Y") if e.date_fin else "Présent"
                    ),
                    "description": e.description or "",
                    "competences": [c.nom for c in getattr(e, "competences", [])],
                }
                for e in experiences
            ]
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des expériences: {e}")
            return []

    def _get_formations(self) -> List[Dict]:
        """Récupère les formations"""
        try:
            formations = (
                Formation.query.filter_by(user_id=self.etudiant.user_id)
                .order_by(Formation.date_debut.desc())
                .all()
            )

            return [
                {
                    "id": f.id,
                    "diplome": f.diplome or "",
                    "etablissement": f.etablissement or "",
                    "lieu": "",  # Corrigé: f.lieu n'existe pas dans le modèle Formation
                    "date_debut": (
                        f.date_debut.strftime("%m/%Y") if f.date_debut else ""
                    ),
                    "date_fin": (
                        f.date_fin.strftime("%m/%Y") if f.date_fin else "Présent"
                    ),
                    "description": f.description or "",
                }
                for f in formations
            ]
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des formations: {e}")
            return []

    def _get_langues(self) -> List[Dict]:
        """Récupère les langues"""
        try:
            langues = Langue.query.filter_by(user_id=self.etudiant.user_id).all()
            return [
                {
                    "id": l.id,
                    "nom": l.nom,
                    "niveau": f"{l.niveau_ecrit or ''} / {l.niveau_oral or ''}",  # Corrigé: l.niveau -> niveaux écrit/oral
                    "certification": getattr(l, "certification", None),
                }
                for l in langues  # noqa
            ]
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des langues: {e}")
            return []

    def _get_certifications(self) -> List[Dict]:
        """Récupère les certifications (si le modèle existe)"""
        try:
            # Vérifier si le modèle Certification existe
            try:
                from app.models.certification import Certification
                certifications = Certification.query.filter_by(user_id=self.etudiant.user_id).all()
                return [
                    {
                        "id": c.id,
                        "nom": c.nom,
                        "organisme": c.organisme,
                        "date_obtention": c.date_obtention.strftime("%m/%Y") if c.date_obtention else "",
                        "date_expiration": c.date_expiration.strftime("%m/%Y") if c.date_expiration else "",
                        "description": c.description or "",
                    }
                    for c in certifications
                ]
            except ImportError:
                # Le modèle n'existe pas, retourner une liste vide
                return []
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des certifications: {e}")
            return []

    def _get_references(self) -> List[Dict]:
        """Récupère les références (si le modèle existe)"""
        try:
            # Vérifier si le modèle Reference existe
            try:
                from app.models.reference import Reference
                references = Reference.query.filter_by(user_id=self.etudiant.user_id).all()
                return [
                    {
                        "id": r.id,
                        "nom": r.nom,
                        "poste": r.poste,
                        "entreprise": r.entreprise,
                        "email": r.email,
                        "telephone": r.telephone,
                        "relation": r.relation,
                    }
                    for r in references
                ]
            except ImportError:
                # Le modèle n'existe pas, retourner une liste vide
                return []
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des références: {e}")
            return []

    def _get_centres_interet(self) -> List[Dict]:
        """Récupère les centres d'intérêt (si le modèle existe)"""
        try:
            # Vérifier si le modele CentreInteret existe
            try:
                from app.models.centre_interet import CentreInteret
                centres = CentreInteret.query.filter_by(user_id=self.etudiant.user_id).all()
                return [
                    {
                        "id": c.id,
                        "nom": c.nom,
                        "description": c.description or "",
                        "categorie": c.categorie or "Général",
                    }
                    for c in centres
                ]
            except ImportError:
                # Le modèle n'existe pas, retourner une liste vide
                return []
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des centres d'intérêt: {e}")
            return []

    def _get_projets(self) -> List[Dict]:
        try:
            projets = Projet.query.filter_by(user_id=self.etudiant.user_id).all()
            return [
                {
                    "id": p.id,
                    "titre": p.titre or "",
                    "description": p.description or "",
                    "lien": getattr(p, "lien", "") or "",
                    "technologies": (
                        p.technologies.split(",") if p.technologies else []
                    ),
                    "date_debut": (
                        p.date_debut.strftime("%m/%Y") if p.date_debut else ""
                    ),
                    "date_fin": (
                        p.date_fin.strftime("%m/%Y") if p.date_fin else "Présent"
                    ),
                    "en_cours": p.en_cours,
                }
                for p in projets
            ]
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des projets: {e}")
            return []

    def _get_skill_level(self, note: float) -> str:
        """Détermine le niveau de compétence basé sur la note"""
        for (min_note, max_note), niveau in self.SKILL_LEVELS.items():
            if min_note <= note < max_note:
                return niveau
        return "Débutant"

    def _generate_content_with_gemini(self, section: str, data: Dict) -> str:
        """Génère le contenu d'une section du CV avec Gemini"""
        if not self.gemini:
            logger.warning("Gemini non disponible, utilisation du format par défaut")
            return self._format_fallback(data)

        prompt = self._build_gemini_prompt(section, data)

        try:
            response = self.gemini.generate_response(prompt=prompt, temperature=0.3)
            return response.get("text", self._format_fallback(data))
        except Exception as e:
            logger.error(f"Erreur Gemini pour {section}: {e}")
            return self._format_fallback(data)

    def _build_gemini_prompt(self, section: str, data: Dict) -> str:
        """Construit le prompt pour Gemini"""
        return f"""
Tu es un expert en rédaction de CV professionnel.

Section: {section}
Données: {json.dumps(data, indent=2, ensure_ascii=False)}

Instructions:
1. Rédaction concise et impactante
2. Phrases courtes avec verbes d'action
3. Résultats mesurables quand possible
4. Structure claire avec puces
5. Ton professionnel et positif
6. Temps verbal cohérent

Génère uniquement le texte formaté, sans explications additionnelles.
"""

    def _format_fallback(self, data: Dict) -> str:
        """Format de secours si Gemini n'est pas disponible"""
        if isinstance(data, dict):
            return "\n".join([f"• {k}: {v}" for k, v in data.items() if v])
        return str(data)

    def generate_pdf(self, template_id: str = "modern") -> Optional[str]:
        """Génère un CV au format PDF"""
        user_data = self.get_user_data()
        if "error" in user_data:
            logger.error(f"Impossible de générer le PDF: {user_data['error']}")
            return None

        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = (
                f"CV_{user_data['utilisateur']['prenom']}_"
                f"{user_data['utilisateur']['nom']}_{timestamp}.pdf"
            )
            output_path = OUTPUT_FOLDER / filename

            # Création du document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            story = self._build_pdf_content(user_data)
            doc.build(story)

            logger.info(f"PDF généré avec succès: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {e}", exc_info=True)
            return None

    def _build_pdf_content(self, user_data: Dict) -> List:
        """Construit le contenu du PDF"""
        styles = getSampleStyleSheet()
        self._enhance_pdf_styles(styles)

        story = []

        # En-tête
        story.extend(self._add_pdf_header(user_data, styles))

        # Sections
        if user_data["utilisateur"]["bio"]:
            story.extend(
                self._add_pdf_section("PROFIL", user_data["utilisateur"]["bio"], styles)
            )

        if user_data["experiences"]:
            story.extend(self._add_pdf_experiences(user_data["experiences"], styles))

        if user_data["formations"]:
            story.extend(self._add_pdf_formations(user_data["formations"], styles))

        if user_data["competences_techniques"]:
            story.extend(
                self._add_pdf_technical_skills(
                    user_data["competences_techniques"], styles
                )
            )

        if user_data["competences"]:
            story.extend(self._add_pdf_skills(user_data["competences"], styles))

        if user_data["langues"]:
            story.extend(self._add_pdf_languages(user_data["langues"], styles))

        if user_data["projets"]:
            story.extend(self._add_pdf_projects(user_data["projets"], styles))

        return story

    def _enhance_pdf_styles(self, styles):
        """Améliore les styles PDF"""
        styles.add(
            ParagraphStyle(
                name="CVTitle",
                parent=styles["Heading1"],
                fontSize=20,
                leading=24,
                textColor=colors.HexColor("#2c3e50"),
                spaceAfter=12,
                fontName="Helvetica-Bold",
            )
        )

        styles.add(
            ParagraphStyle(
                name="CVSection",
                parent=styles["Heading2"],
                fontSize=14,
                leading=18,
                textColor=colors.HexColor("#34495e"),
                spaceAfter=10,
                spaceBefore=15,
                fontName="Helvetica-Bold",
            )
        )

    def _add_pdf_header(self, user_data: Dict, styles) -> List:
        """Ajoute l'en-tête du PDF"""
        story = []
        user = user_data["utilisateur"]

        title = f"{user['prenom']} {user['nom']}"
        story.append(Paragraph(title, styles["CVTitle"]))

        contact_info = [
            user["email"],
            user["telephone"],
            f"{user['adresse']}, {user['code_postal']} {user['ville']}".strip(", "),
            user["linkedin"],
            user["github"],
        ]
        contact_info = [info for info in contact_info if info]

        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), styles["Normal"]))
        story.append(Spacer(1, 0.3 * inch))

        return story

    def _add_pdf_section(self, title: str, content: str, styles) -> List:
        """Ajoute une section au PDF"""
        story = []
        story.append(Paragraph(title, styles["CVSection"]))
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))
        return story

    def _add_pdf_experiences(self, experiences: List[Dict], styles) -> List:
        """Ajoute les expériences au PDF"""
        story = []
        story.append(Paragraph("EXPÉRIENCES PROFESSIONNELLES", styles["CVSection"]))

        for exp in experiences:
            title = f"<b>{exp['titre']}</b> chez {exp['entreprise']}"
            story.append(Paragraph(title, styles["Normal"]))

            dates = f"<i>{exp['date_debut']} - {exp['date_fin']} | {exp['lieu']}</i>"
            story.append(Paragraph(dates, styles["Normal"]))

            if exp["description"]:
                story.append(Paragraph(exp["description"], styles["Normal"]))

            story.append(Spacer(1, 0.15 * inch))

        return story

    def _add_pdf_formations(self, formations: List[Dict], styles) -> List:
        """Ajoute les formations au PDF"""
        story = []
        story.append(Paragraph("FORMATIONS", styles["CVSection"]))

        for form in formations:
            title = f"<b>{form['diplome']}</b> - {form['etablissement']}"
            story.append(Paragraph(title, styles["Normal"]))

            dates = f"<i>{form['date_debut']} - {form['date_fin']} | {form['lieu']}</i>"
            story.append(Paragraph(dates, styles["Normal"]))

            if form["description"]:
                story.append(Paragraph(form["description"], styles["Normal"]))

            story.append(Spacer(1, 0.15 * inch))

        return story

    def _add_pdf_technical_skills(self, skills: List[Dict], styles) -> List:
        """Ajoute les compétences techniques au PDF"""
        story = []
        story.append(Paragraph("COMPÉTENCES TECHNIQUES", styles["CVSection"]))

        # Regroupement par catégorie
        skills_by_category = {}
        for skill in skills:
            category = skill.get("categorie", "Autre")
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(f"{skill['nom']} ({skill['niveau']})")

        for category, skills_list in skills_by_category.items():
            content = f"<b>{category}:</b> {', '.join(skills_list)}"
            story.append(Paragraph(content, styles["Normal"]))

        story.append(Spacer(1, 0.2 * inch))
        return story

    def _add_pdf_skills(self, skills: List[Dict], styles) -> List:
        """Ajoute les compétences générales au PDF"""
        story = []
        story.append(Paragraph("COMPÉTENCES GÉNÉRALES", styles["CVSection"]))

        skills_text = ", ".join([f"{s['nom']} ({s['niveau']})" for s in skills])
        story.append(Paragraph(skills_text, styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))

        return story

    def _add_pdf_languages(self, languages: List[Dict], styles) -> List:
        """Ajoute les langues au PDF"""
        story = []
        story.append(Paragraph("LANGUES", styles["CVSection"]))

        lang_items = []
        for lang in languages:
            cert = f" ({lang['certification']})" if lang.get("certification") else ""
            lang_items.append(f"{lang['nom']} - {lang['niveau']}{cert}")

        story.append(Paragraph(" • ".join(lang_items), styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))

        return story

    def _add_pdf_projects(self, projects: List[Dict], styles) -> List:
        """Ajoute les projets au PDF"""
        story = []
        story.append(Paragraph("PROJETS", styles["CVSection"]))

        for proj in projects:
            story.append(Paragraph(f"<b>{proj['titre']}</b>", styles["Normal"]))

            if proj.get("date_realisation"):
                story.append(
                    Paragraph(
                        f"<i>Réalisé en {proj['date_realisation']}</i>",
                        styles["Normal"],
                    )
                )

            if proj["description"]:
                story.append(Paragraph(proj["description"], styles["Normal"]))

            if proj.get("technologies"):
                tech_text = f"<b>Technologies:</b> {', '.join(proj['technologies'])}"
                story.append(Paragraph(tech_text, styles["Normal"]))

            if proj.get("lien"):
                story.append(
                    Paragraph(f"<b>Lien:</b> {proj['lien']}", styles["Normal"])
                )

            story.append(Spacer(1, 0.15 * inch))

        return story

    def generate_docx(self, template_id: str = "modern") -> Optional[str]:
        """Génère un CV au format Word (DOCX)"""
        user_data = self.get_user_data()
        if "error" in user_data:
            logger.error(f"Impossible de générer le DOCX: {user_data['error']}")
            return None

        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = (
                f"CV_{user_data['utilisateur']['prenom']}_"
                f"{user_data['utilisateur']['nom']}_{timestamp}.docx"
            )
            output_path = OUTPUT_FOLDER / filename

            doc = Document()
            self._configure_docx_styles(doc)
            self._build_docx_content(doc, user_data)

            doc.save(str(output_path))
            logger.info(f"DOCX généré avec succès: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Erreur lors de la génération du DOCX: {e}", exc_info=True)
            return None

    def _configure_docx_styles(self, doc: Document):
        """Configure les styles du document Word"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _build_docx_content(self, doc: Document, user_data: Dict):
        """Construit le contenu du document Word"""
        user = user_data["utilisateur"]

        # En-tête
        title = doc.add_heading(level=1)
        title_run = title.add_run(f"{user['prenom']} {user['nom']}")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(44, 62, 80)

        # Contact
        contact_info = [
            user["email"],
            user["telephone"],
            f"{user['adresse']}, {user['code_postal']} {user['ville']}".strip(", "),
            user["linkedin"],
            user["github"],
        ]
        contact_info = [info for info in contact_info if info]

        if contact_info:
            coord_para = doc.add_paragraph(" | ".join(contact_info))
            coord_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()

        # Profil
        if user["bio"]:
            doc.add_heading("PROFIL", level=2)
            doc.add_paragraph(user["bio"])
            doc.add_paragraph()

        # Expériences
        if user_data["experiences"]:
            doc.add_heading("EXPÉRIENCES PROFESSIONNELLES", level=2)
            for exp in user_data["experiences"]:
                self._add_docx_experience(doc, exp)

        # Formations
        if user_data["formations"]:
            doc.add_heading("FORMATIONS", level=2)
            for form in user_data["formations"]:
                self._add_docx_formation(doc, form)

        # Compétences techniques
        if user_data["competences_techniques"]:
            self._add_docx_technical_skills(doc, user_data["competences_techniques"])

        # Compétences générales
        if user_data["competences"]:
            self._add_docx_skills(doc, user_data["competences"])

        # Langues
        if user_data["langues"]:
            self._add_docx_languages(doc, user_data["langues"])

        # Projets
        if user_data["projets"]:
            doc.add_heading("PROJETS", level=2)
            for proj in user_data["projets"]:
                self._add_docx_project(doc, proj)

    def _add_docx_experience(self, doc: Document, exp: Dict):
        """Ajoute une expérience au document Word"""
        exp_title = doc.add_heading(level=3)
        exp_title_run = exp_title.add_run(f"{exp['titre']} chez {exp['entreprise']}")
        exp_title_run.bold = True

        exp_dates = doc.add_paragraph()
        exp_dates_run = exp_dates.add_run(
            f"{exp['date_debut']} - {exp['date_fin']} | {exp['lieu']}"
        )
        exp_dates_run.italic = True

        if exp["description"]:
            doc.add_paragraph(exp["description"])

        if exp.get("competences"):
            comp_para = doc.add_paragraph()
            comp_run = comp_para.add_run("Compétences: ")
            comp_run.bold = True
            comp_para.add_run(", ".join(exp["competences"]))

        doc.add_paragraph()

    def _add_docx_formation(self, doc: Document, form: Dict):
        """Ajoute une formation au document Word"""
        form_title = doc.add_heading(level=3)
        form_title_run = form_title.add_run(
            f"{form['diplome']} - {form['etablissement']}"
        )
        form_title_run.bold = True

        form_dates = doc.add_paragraph()
        form_dates_run = form_dates.add_run(
            f"{form['date_debut']} - {form['date_fin']} | {form['lieu']}"
        )
        form_dates_run.italic = True

        if form["description"]:
            doc.add_paragraph(form["description"])

        doc.add_paragraph()

    def _add_docx_technical_skills(self, doc: Document, skills: List[Dict]):
        """Ajoute les compétences techniques au document Word"""
        doc.add_heading("COMPÉTENCES TECHNIQUES", level=2)

        skills_by_category = {}
        for skill in skills:
            category = skill.get("categorie", "Autre")
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(f"{skill['nom']} ({skill['niveau']})")

        for category, skills_list in skills_by_category.items():
            cat_para = doc.add_paragraph()
            cat_run = cat_para.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.underline = True
            cat_para.add_run(", ".join(skills_list))

        doc.add_paragraph()

    def _add_docx_skills(self, doc: Document, skills: List[Dict]):
        """Ajoute les compétences générales au document Word"""
        doc.add_heading("COMPÉTENCES GÉNÉRALES", level=2)
        skills_text = ", ".join([f"{s['nom']} ({s['niveau']})" for s in skills])
        doc.add_paragraph(skills_text)
        doc.add_paragraph()

    def _add_docx_languages(self, doc: Document, languages: List[Dict]):
        """Ajoute les langues au document Word"""
        doc.add_heading("LANGUES", level=2)
        lang_items = []
        for lang in languages:
            cert = f" ({lang['certification']})" if lang.get("certification") else ""
            lang_items.append(f"{lang['nom']} - {lang['niveau']}{cert}")
        doc.add_paragraph(" • ".join(lang_items))
        doc.add_paragraph()

    def _add_docx_project(self, doc: Document, proj: Dict):
        """Ajoute un projet au document Word"""
        proj_title = doc.add_heading(level=3)
        proj_title_run = proj_title.add_run(proj["titre"])
        proj_title_run.bold = True

        if proj.get("date_realisation"):
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Réalisé en {proj['date_realisation']}")
            date_run.italic = True

        if proj["description"]:
            doc.add_paragraph(proj["description"])

        if proj.get("technologies"):
            tech_para = doc.add_paragraph()
            tech_run = tech_para.add_run("Technologies: ")
            tech_run.bold = True
            tech_para.add_run(", ".join(proj["technologies"]))

        if proj.get("lien"):
            link_para = doc.add_paragraph()
            link_run = link_para.add_run("Lien: ")
            link_run.bold = True
            link_para.add_run(proj["lien"])

        doc.add_paragraph()

    def generate_preview(self, template_id: str = "modern") -> str:
        """Génère un aperçu HTML du CV en utilisant Gemini"""
        try:
            logger.info(f"Génération de l'aperçu pour l'utilisateur {self.user_id}")

            user_data = self.get_user_data()

            # Gestion des erreurs
            if "error" in user_data:
                return self._generate_error_html(user_data["error"])

            if not user_data or not user_data.get("utilisateur"):
                return self._generate_error_html(
                    "Données utilisateur manquantes. Veuillez compléter votre profil."
                )

            logger.info(f"Données récupérées: {', '.join(user_data.keys())}")

            # Si Gemini n'est pas disponible, générer un aperçu basique
            if not self.gemini:
                logger.warning("Gemini non disponible, génération d'un aperçu basique")
                return self._generate_basic_html_preview(user_data)

            # Construire le prompt pour Gemini
            prompt = self._build_preview_prompt(user_data)

            logger.info("Appel à l'API Gemini pour l'aperçu...")
            response = self.gemini.generate_response(prompt)

            if not response or not isinstance(response, dict):
                logger.error("Réponse invalide de Gemini")
                return self._generate_basic_html_preview(user_data)

            if not response.get("success"):
                error_msg = response.get("error", "Erreur inconnue")
                logger.error(f"Erreur Gemini: {error_msg}")
                return self._generate_basic_html_preview(user_data)

            html_content = response.get("response", "").strip()

            if not html_content:
                logger.warning("Contenu HTML vide, utilisation de l'aperçu basique")
                return self._generate_basic_html_preview(user_data)

            # Nettoyer le contenu HTML
            html_content = self._clean_html_content(html_content)

            # Encapsuler dans un document HTML complet si nécessaire
            if "<!DOCTYPE html>" not in html_content[:100]:
                html_content = self._wrap_in_html_document(html_content, user_data)

            logger.info("Aperçu généré avec succès")
            return html_content

        except Exception as e:
            logger.error(f"Erreur dans generate_preview: {e}", exc_info=True)
            return self._generate_error_html(str(e))

    def _build_preview_prompt(self, user_data: Dict) -> str:
        """Construit le prompt pour la génération de l'aperçu"""
        # Créer une copie des données utilisateur pour les modifier
        safe_user_data = user_data.copy()

        # Simplifier les données pour réduire la taille du prompt
        if "experiences" in safe_user_data and len(safe_user_data["experiences"]) > 3:
            safe_user_data["experiences"] = safe_user_data["experiences"][:3]
            safe_user_data["_note"] = (
                "... (expériences supplémentaires non affichées pour économiser des tokens)"
            )

        if "formations" in safe_user_data and len(safe_user_data["formations"]) > 2:
            safe_user_data["formations"] = safe_user_data["formations"][:2]

        if (
            "competences_techniques" in safe_user_data
            and len(safe_user_data["competences_techniques"]) > 10
        ):
            safe_user_data["competences_techniques"] = safe_user_data[
                "competences_techniques"
            ][:10]

        # Créer un prompt plus court et plus directif
        prompt = """
GÉNÈRE UNE PAGE HTML DE CV avec ces caractéristiques :
- Utilise Tailwind CSS (classes comme bg-blue-100, p-4, rounded-lg, etc.)
- Structure : En-tête (nom, contact, photo), Profil, Expériences, Formations, Compétences, Langues, Projets, Certifications, Références, Centres d'intérêt
- Design moderne, épuré et responsive
- Couleurs : bleus et gris professionnels
- Sections en cartes avec ombres légères
- Pas de commentaires, juste le HTML

EXEMPLE DE STRUCTURE :
```
<div class="max-w-4xl mx-auto p-4">
  <header class="text-center mb-8">
    <div class="flex items-center justify-center mb-4">
      <img src="photo_url" class="w-20 h-20 rounded-full mr-4" alt="Photo">
      <div>
        <h1 class="text-3xl font-bold">Prénom NOM</h1>
        <p class="text-gray-600">Titre professionnel</p>
        <div class="mt-2 space-x-4">
          <span>email@exemple.com</span>
          <span>•</span>
          <span>+33 6 12 34 56 78</span>
        </div>
        <div class="mt-1 text-sm text-gray-500">
          <span>Disponibilité: Immédiate</span>
          <span>•</span>
          <span>Permis: B</span>
        </div>
      </div>
    </div>
  </header>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Profil</h2>
    <p class="mt-2">Bio/résumé professionnel concis.</p>
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Expériences</h2>
    <!-- Expériences avec dates et descriptions -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Formations</h2>
    <!-- Formations avec diplômes et établissements -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Compétences</h2>
    <!-- Compétences techniques et générales -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Langues</h2>
    <!-- Langues avec niveaux -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Projets</h2>
    <!-- Projets pertinents -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Certifications</h2>
    <!-- Certifications professionnelles -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Références</h2>
    <!-- Références professionnelles -->
  </section>
  
  <section class="mb-8">
    <h2 class="text-xl font-semibold border-b-2 border-blue-500 pb-1">Centres d'intérêt</h2>
    <!-- Loisirs et activités -->
  </section>
</div>
```

DONNÉES À UTILISER (format JSON simplifié) :
"""

        # Ajouter les données utilisateur simplifiées
        prompt += json.dumps(safe_user_data, indent=2, ensure_ascii=False, default=str)

        # Instructions finales
        prompt += """

INSTRUCTIONS FINALES :
1. Génère UNIQUEMENT le code HTML (sans balises ```html)
2. Utilise des classes Tailwind pour le style
3. Sois concis mais professionnel
4. Inclus toutes les sections pertinentes des données fournies :
   - En-tête avec photo, titre professionnel, disponibilité, permis
   - Profil (bio)
   - Expériences professionnelles
   - Formations
   - Compétences (techniques et générales)
   - Langues avec niveaux
   - Projets pertinents
   - Certifications (si disponibles)
   - Références (si disponibles)
   - Centres d'intérêt (si disponibles)
5. Formatte les dates de manière lisible (ex: "Janvier 2020 - Mars 2022")
6. Assure-toi que le code est valide et bien structuré
7. N'affiche que les sections qui ont des données (sauf Profil qui est obligatoire)
8. Pour la photo, utilise l'URL fournie ou une icône par défaut si vide
"""

        return prompt

    def _clean_html_content(self, html_content: str) -> str:
        """Nettoie le contenu HTML reçu"""
        # Supprimer les balises de code Markdown
        html_content = re.sub(r"```html\n?", "", html_content)
        html_content = re.sub(r"```\n?", "", html_content)

        # Extraire le HTML si enveloppé dans du texte
        if not html_content.strip().startswith("<"):
            html_match = re.search(r"(<[^>]+>.*</[^>]+>)", html_content, re.DOTALL)
            if html_match:
                html_content = html_match.group(1)

        return html_content.strip()

    def _wrap_in_html_document(self, content: str, user_data: Dict) -> str:
        """Encapsule le contenu dans un document HTML complet"""
        user = user_data.get("utilisateur", {})
        title = f"CV - {user.get('prenom', '')} {user.get('nom', '')}".strip()

        return f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>
        @media print {{
            body {{ margin: 0; padding: 20px; }}
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body class="bg-gray-50 min-h-screen">
    <div class="container mx-auto p-4 max-w-4xl">
        {content}
    </div>
</body>
</html>"""

    def _generate_basic_html_preview(self, user_data: Dict) -> str:
        """Génère un aperçu HTML basique sans Gemini"""
        user = user_data.get("utilisateur", {})

        html_parts = [
            "<!DOCTYPE html>",
            '<html lang="fr">',
            "<head>",
            '    <meta charset="UTF-8">',
            '    <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f'    <title>CV - {user.get("prenom", "")} {user.get("nom", "")}</title>',
            '    <script src="https://cdn.tailwindcss.com"></script>',
            "</head>",
            '<body class="bg-gray-50 p-8">',
            '    <div class="max-w-4xl mx-auto bg-white rounded-lg shadow-lg p-8">',
        ]

        # En-tête
        html_parts.extend(
            [
                f'        <h1 class="text-4xl font-bold text-gray-800 mb-2">{user.get("prenom", "")} {user.get("nom", "")}</h1>',
                '        <div class="text-gray-600 mb-6">',
            ]
        )

        contact_items = [
            user.get("email"),
            user.get("telephone"),
            f"{user.get('ville', '')} {user.get('code_postal', '')}".strip(),
        ]
        contact_items = [item for item in contact_items if item]

        if contact_items:
            html_parts.append(f'            <p>{" | ".join(contact_items)}</p>')

        html_parts.append("        </div>")

        # Profil
        if user.get("bio"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Profil</h2>',
                    f'            <p class="text-gray-700">{user["bio"]}</p>',
                    "        </section>",
                ]
            )

        # Expériences
        if user_data.get("experiences"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Expériences Professionnelles</h2>',
                ]
            )

            for exp in user_data["experiences"]:
                html_parts.extend(
                    [
                        '            <div class="mb-4 p-4 bg-gray-50 rounded">',
                        f'                <h3 class="text-xl font-semibold text-gray-800">{exp.get("titre", "")}</h3>',
                        f'                <p class="text-gray-600">{exp.get("entreprise", "")} | {exp.get("lieu", "")}</p>',
                        f'                <p class="text-sm text-gray-500">{exp.get("date_debut", "")} - {exp.get("date_fin", "")}</p>',
                    ]
                )

                if exp.get("description"):
                    html_parts.append(
                        f'                <p class="mt-2 text-gray-700">{exp["description"]}</p>'
                    )

                html_parts.append("            </div>")

            html_parts.append("        </section>")

        # Formations
        if user_data.get("formations"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Formations</h2>',
                ]
            )

            for form in user_data["formations"]:
                html_parts.extend(
                    [
                        '            <div class="mb-4 p-4 bg-gray-50 rounded">',
                        f'                <h3 class="text-xl font-semibold text-gray-800">{form.get("diplome", "")}</h3>',
                        f'                <p class="text-gray-600">{form.get("etablissement", "")} | {form.get("lieu", "")}</p>',
                        f'                <p class="text-sm text-gray-500">{form.get("date_debut", "")} - {form.get("date_fin", "")}</p>',
                    ]
                )

                if form.get("description"):
                    html_parts.append(
                        f'                <p class="mt-2 text-gray-700">{form["description"]}</p>'
                    )

                html_parts.append("            </div>")

            html_parts.append("        </section>")

        # Compétences
        if user_data.get("competences") or user_data.get("competences_techniques"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Compétences</h2>',
                    '            <div class="flex flex-wrap gap-2">',
                ]
            )

            all_skills = user_data.get("competences", []) + user_data.get(
                "competences_techniques", []
            )

            for skill in all_skills:
                html_parts.append(
                    f'                <span class="px-3 py-1 bg-blue-100 text-blue-800 rounded-full text-sm">'
                    f'{skill.get("nom", "")} - {skill.get("niveau", "")}</span>'
                )

            html_parts.extend(
                [
                    "            </div>",
                    "        </section>",
                ]
            )

        # Langues
        if user_data.get("langues"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Langues</h2>',
                    '            <div class="flex flex-wrap gap-2">',
                ]
            )

            for lang in user_data["langues"]:
                html_parts.append(
                    f'                <span class="px-3 py-1 bg-green-100 text-green-800 rounded-full text-sm">'
                    f'{lang.get("nom", "")} - {lang.get("niveau", "")}</span>'
                )

            html_parts.extend(
                [
                    "            </div>",
                    "        </section>",
                ]
            )

        # Projets
        if user_data.get("projets"):
            html_parts.extend(
                [
                    '        <section class="mb-8">',
                    '            <h2 class="text-2xl font-bold text-gray-800 mb-3 border-b-2 border-blue-500 pb-2">Projets</h2>',
                ]
            )

            for proj in user_data["projets"]:
                html_parts.extend(
                    [
                        '            <div class="mb-4 p-4 bg-gray-50 rounded">',
                        f'                <h3 class="text-xl font-semibold text-gray-800">{proj.get("titre", "")}</h3>',
                    ]
                )

                if proj.get("description"):
                    html_parts.append(
                        f'                <p class="mt-2 text-gray-700">{proj["description"]}</p>'
                    )

                if proj.get("technologies"):
                    tech_badges = " ".join(
                        [
                            f'<span class="px-2 py-1 bg-purple-100 text-purple-800 rounded text-xs">{tech}</span>'
                            for tech in proj["technologies"]
                        ]
                    )
                    html_parts.append(
                        f'                <div class="mt-2 flex flex-wrap gap-1">{tech_badges}</div>'
                    )

                if proj.get("lien"):
                    html_parts.append(
                        f'                <p class="mt-2"><a href="{proj["lien"]}" class="text-blue-600 hover:underline" target="_blank">'
                        f"Voir le projet</a></p>"
                    )

                html_parts.append("            </div>")

            html_parts.append("        </section>")

        # Fermeture
        html_parts.extend(
            [
                "    </div>",
                "</body>",
                "</html>",
            ]
        )

        return "\n".join(html_parts)

    def _generate_error_html(self, error_message: str) -> str:
        """Génère un message d'erreur HTML"""
        return f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Erreur - Génération CV</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-50 p-8">
    <div class="max-w-2xl mx-auto bg-white rounded-lg shadow-lg p-8">
        <div class="flex items-center mb-4">
            <svg class="w-12 h-12 text-red-500 mr-4" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                      d="M12 8v4m0 4h.01M21 12a9 9 0 11-18 0 9 9 0 0118 0z"></path>
            </svg>
            <h1 class="text-2xl font-bold text-gray-800">Erreur de génération</h1>
        </div>
        <div class="bg-red-50 border-l-4 border-red-500 p-4 mb-4">
            <p class="text-red-700">{error_message}</p>
        </div>
        <div class="mt-6">
            <h2 class="font-semibold text-gray-800 mb-2">Actions suggérées:</h2>
            <ul class="list-disc list-inside text-gray-600 space-y-1">
                <li>Vérifiez que votre profil est complet</li>
                <li>Assurez-vous d'avoir un profil étudiant actif</li>
                <li>Contactez le support si le problème persiste</li>
            </ul>
        </div>
    </div>
</body>
</html>"""

    def generate_linkedin_data(self) -> Dict[str, Any]:
        """Génère les données au format LinkedIn"""
        user_data = self.get_user_data()
        if "error" in user_data:
            logger.error(
                f"Impossible de générer les données LinkedIn: {user_data['error']}"
            )
            return {"error": user_data["error"]}

        try:
            user = user_data["utilisateur"]

            linkedin_data = {
                "profile": {
                    "firstName": user.get("prenom", ""),
                    "lastName": user.get("nom", ""),
                    "headline": self._generate_headline(user_data),
                    "summary": user.get("bio", ""),
                    "location": self._build_location(user),
                    "contactInfo": {
                        "email": user.get("email", ""),
                        "phone": user.get("telephone", ""),
                        "address": user.get("adresse", ""),
                    },
                    "websites": self._build_websites(user),
                },
                "experience": self._format_linkedin_experiences(
                    user_data.get("experiences", [])
                ),
                "education": self._format_linkedin_education(
                    user_data.get("formations", [])
                ),
                "skills": self._format_linkedin_skills(user_data),
                "languages": self._format_linkedin_languages(
                    user_data.get("langues", [])
                ),
                "projects": self._format_linkedin_projects(
                    user_data.get("projets", [])
                ),
            }

            logger.info("Données LinkedIn générées avec succès")
            return linkedin_data

        except Exception as e:
            logger.error(
                f"Erreur lors de la génération des données LinkedIn: {e}", exc_info=True
            )
            return {"error": str(e)}

    def _generate_headline(self, user_data: Dict) -> str:
        """Génère un titre professionnel pour LinkedIn"""
        formations = user_data.get("formations", [])
        if formations:
            return f"Étudiant en {formations[0].get('diplome', '')}"

        experiences = user_data.get("experiences", [])
        if experiences:
            return experiences[0].get("titre", "")

        return "Professionnel"

    def _build_location(self, user: Dict) -> Dict[str, str]:
        """Construit l'objet de localisation"""
        location = {}
        if user.get("pays"):
            location["country"] = user["pays"]
        if user.get("code_postal"):
            location["postalCode"] = user["code_postal"]
        if user.get("ville"):
            location["city"] = user["ville"]
        return location

    def _build_websites(self, user: Dict) -> List[Dict[str, str]]:
        """Construit la liste des sites web"""
        websites = []
        if user.get("linkedin"):
            websites.append({"type": "LinkedIn", "url": user["linkedin"]})
        if user.get("github"):
            websites.append({"type": "GitHub", "url": user["github"]})
        return websites

    def _format_linkedin_experiences(self, experiences: List[Dict]) -> List[Dict]:
        """Formate les expériences pour LinkedIn"""
        formatted = []
        for exp in experiences:
            formatted.append(
                {
                    "title": exp.get("titre", ""),
                    "company": exp.get("entreprise", ""),
                    "location": exp.get("lieu", ""),
                    "description": exp.get("description", ""),
                    "startDate": self._parse_date(exp.get("date_debut", "")),
                    "endDate": self._parse_date(exp.get("date_fin", "")),
                    "current": exp.get("date_fin") == "Présent",
                }
            )
        return formatted

    def _format_linkedin_education(self, formations: List[Dict]) -> List[Dict]:
        """Formate les formations pour LinkedIn"""
        formatted = []
        for edu in formations:
            formatted.append(
                {
                    "school": edu.get("etablissement", ""),
                    "degree": edu.get("diplome", ""),
                    "fieldOfStudy": edu.get("domaine", ""),
                    "startDate": self._parse_date(edu.get("date_debut", "")),
                    "endDate": self._parse_date(edu.get("date_fin", "")),
                    "description": edu.get("description", ""),
                }
            )
        return formatted

    def _format_linkedin_skills(self, user_data: Dict) -> List[Dict]:
        """Formate les compétences pour LinkedIn"""
        skills = []

        for comp in user_data.get("competences", []):
            skills.append(
                {
                    "name": comp.get("nom", ""),
                    "level": comp.get("niveau", ""),
                    "category": comp.get("categorie", "Général"),
                }
            )

        for comp_tech in user_data.get("competences_techniques", []):
            skills.append(
                {
                    "name": comp_tech.get("nom", ""),
                    "level": comp_tech.get("niveau", ""),
                    "category": comp_tech.get("categorie", "Technique"),
                }
            )

        return skills

    def _format_linkedin_languages(self, langues: List[Dict]) -> List[Dict]:
        """Formate les langues pour LinkedIn"""
        return [
            {
                "name": lang.get("nom", ""),
                "proficiency": lang.get("niveau", ""),
                "certification": lang.get("certification", ""),
            }
            for lang in langues
        ]

    def _format_linkedin_projects(self, projets: List[Dict]) -> List[Dict]:
        """Formate les projets pour LinkedIn"""
        formatted = []
        for proj in projets:
            formatted.append(
                {
                    "name": proj.get("titre", ""),
                    "description": proj.get("description", ""),
                    "url": proj.get("lien", ""),
                    "startDate": self._parse_date(proj.get("date_realisation", "")),
                    "technologies": proj.get("technologies", []),
                }
            )
        return formatted

    def _parse_date(self, date_str: str) -> Optional[Dict[str, int]]:
        """Parse une date au format MM/YYYY"""
        if not date_str or date_str == "Présent":
            return None

        try:
            parts = date_str.split("/")
            if len(parts) == 2:
                return {"month": int(parts[0]), "year": int(parts[1])}
        except (ValueError, IndexError):
            logger.warning(f"Impossible de parser la date: {date_str}")

        return None


# ============================================================================
# ROUTES API
# ============================================================================


@career_craft_bp.route("/templates", methods=["GET"])
@login_required
def get_templates():
    """
    Récupère la liste des modèles de CV disponibles.

    Cette fonction est appelée lorsque l'utilisateur souhaite récupérer la liste des
    modèles de CV disponibles. Elle renvoie une réponse contenant une liste de
    dictionnaires représentant chaque modèle de CV. Chaque dictionnaire contient les
    informations suivantes :
    - id : identifiant unique du modèle de CV
    - name : nom du modèle de CV
    - description : description du modèle de CV
    - preview : aperçu du modèle de CV
    - category : catégorie du modèle de CV

    Cette fonction est appelée par une requête GET à l'URL /career_craft/templates.

    Returns:
        Une réponse JSON contenant une liste de dictionnaires. Chaque dictionnaire
        contient les informations suivantes :
        - id : identifiant unique du modèle de CV
        - name : nom du modèle de CV
        - description : description du modèle de CV
        - preview : aperçu du modèle de CV
        - category : catégorie du modèle de CV
    """
    try:
        templates = []

        for file_path in TEMPLATES_FOLDER.glob("*.json"):
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    template_data = json.load(f)
                    templates.append(
                        {
                            "id": template_data.get("id", ""),
                            "name": template_data.get("name", ""),
                            "description": template_data.get("description", ""),
                            "preview": template_data.get("preview", ""),
                            "category": template_data.get("category", "general"),
                        }
                    )
            except Exception as e:
                logger.error(
                    f"Erreur lors du chargement du template {file_path.name}: {e}"
                )

        return jsonify({"success": True, "templates": templates})

    except Exception as e:
        logger.error(f"Erreur dans get_templates: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la récupération des templates: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/generate", methods=["POST"])
@login_required
def generate_cv():
    """
    Génère un CV dans le format demandé.

    Cette fonction est appelée lorsque l'utilisateur souhaite générer son CV.
    Elle prend en entrée une requête POST contenant les informations nécessaires
    au générateur de CV. Ces informations comprennent le format souhaité (pdf,
    docx ou linkedin) et l'ID du modèle de CV à utiliser (par défaut, le modèle
    "modern").

    Si le format n'est pas valide, la fonction renvoie une réponse JSON avec
    un message d'erreur et un code d'état 500.

    Si le format est valide, la fonction démarre le processus de génération du
    CV. Les données nécessaires au générateur de CV sont extraites de la requête
    POST et stockées dans une variable appelée "data". Le format souhaité est extrait
    de "data" et stocké dans une variable appelée "format_type". L'ID du modèle de
    CV est extrait de "data" et stocké dans une variable appelée "template_id".

    Enfin, la fonction renvoie une réponse JSON contenant les informations suivantes :
    - success (bool): Indique si la requête a été effectuée avec succès
    - url (str): Lien vers le fichier généré

    Exemple de réponse JSON :
    {
        "success": true,
        "url": "https://example.com/cv.pdf"
    }
    """
    try:
        data = request.get_json()
        format_type = data.get("format", "pdf").lower()
        template_id = data.get("template", "modern")

        # Validation du format
        valid_formats = ["pdf", "docx", "linkedin"]
        if format_type not in valid_formats:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Format non supporté. Choisissez parmi: {', '.join(valid_formats)}",
                    }
                ),
                400,
            )

        cv_generator = CVGenerator(current_user.id)

        # Génération selon le format
        if format_type == "pdf":
            output_path = cv_generator.generate_pdf(template_id)
        elif format_type == "docx":
            output_path = cv_generator.generate_docx(template_id)
        elif format_type == "linkedin":
            linkedin_data = cv_generator.generate_linkedin_data()
            if "error" in linkedin_data:
                return (
                    jsonify({"success": False, "message": linkedin_data["error"]}),
                    500,
                )

            return jsonify(
                {
                    "success": True,
                    "message": "Données LinkedIn générées avec succès",
                    "data": linkedin_data,
                }
            )

        # Vérification du fichier généré
        if not output_path or not Path(output_path).exists():
            return (
                jsonify(
                    {"success": False, "message": "Erreur lors de la génération du CV"}
                ),
                500,
            )

        # Retourner le chemin relatif pour le téléchargement
        relative_path = Path(output_path).relative_to(current_app.root_path)

        return jsonify(
            {
                "success": True,
                "message": f"CV {format_type.upper()} généré avec succès",
                "download_url": f"/api/career-craft/download?path={relative_path}",
                "filename": Path(output_path).name,
            }
        )

    except Exception as e:
        logger.error(f"Erreur lors de la génération du CV: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la génération du CV: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/download", methods=["GET"])
@login_required
def download_file():
    """
    Télécharge un fichier généré.

    Cette fonction permet à l'utilisateur de télécharger un fichier généré par
    le service de génération de CV. Elle est utilisée pour permettre à l'utilisateur
    de télécharger le CV généré et l'ouvrir directement dans son navigateur.

    Args:
        None

    Returns:
        Une réponse HTTP contenant le fichier généré.

    Raises:
        None

    """
    # Récupérer le chemin relatif du fichier à télécharger depuis les paramètres de la requête
    try:
        file_path = request.args.get("path")

        if not file_path:
            return (
                jsonify({"success": False, "error": "Chemin du fichier manquant"}),
                400,
            )

            # Construire le chemin absolu de manière sécurisée
            absolute_path = Path(current_app.root_path) / file_path

            # Vérifier que le fichier existe
            if not absolute_path.exists():
                logger.error(f"Fichier non trouvé: {absolute_path}")
                return jsonify({"success": False, "error": "Fichier non trouvé"}), 404

            # Vérifier que le fichier est dans un dossier autorisé (sécurité)
            try:
                absolute_path.resolve().relative_to(OUTPUT_FOLDER.resolve())
            except ValueError:
                logger.error(
                    f"Tentative d'accès à un fichier non autorisé: {absolute_path}"
                )
                return jsonify({"success": False, "error": "Accès non autorisé"}), 403

            return send_file(
                str(absolute_path), as_attachment=True, download_name=absolute_path.name
            )

    except Exception as e:
        logger.error(f"Erreur lors du téléchargement: {e}", exc_info=True)
        return (
            jsonify(
                {"success": False, "error": f"Erreur lors du téléchargement: {str(e)}"}
            ),
            500,
        )


@career_craft_bp.route("/preview", methods=["POST"])
@login_required
def preview_cv():
    """
    Génère un aperçu HTML du CV.

    Cette fonction prend en entrée une requête POST contenant les données du CV
    et renvoie une réponse JSON contenant le contenu HTML de l'aperçu du CV.

    Args:
        None

    Returns:
        Une réponse JSON contenant les informations suivantes :
        - success (bool): Indique si la requête a été effectuée avec succès
        - html (str): Contenu HTML de l'aperçu du CV

    Exemple de réponse JSON :
    {
        "success": true,
        "html": "<!DOCTYPE html><html lang="fr"><head><meta charset="UTF-8"><title>CV - Nom Prénom</title></head><body><h1>CV - Nom Prénom</h1></body></html>"
    }
    """
    try:
        data = request.get_json() or {}
        template_id = data.get("template", "modern")

        cv_generator = CVGenerator(current_user.id)
        html_content = cv_generator.generate_preview(template_id)

        return jsonify({"success": True, "html": html_content})

    except Exception as e:
        logger.error(f"Erreur lors de la génération de l'aperçu: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la génération de l'aperçu: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/data", methods=["GET"])
@login_required
def get_user_cv_data():
    """
    Récupère les données de l'utilisateur pour le CV.

    Cette fonction est appelée lorsque l'utilisateur souhaite récupérer les données
    nécessaires pour générer son CV. Elle renvoie une réponse JSON contenant les
    informations suivantes :
    - success (bool): Indique si la requête a été effectuée avec succès
    - data (dict): Données de l'utilisateur (nom, prénom, email, etc.)
    - error (str): Message d'erreur si une erreur est survenue

    Exemple de réponse JSON :
    {
        "success": true,
        "data": {
            "nom": "Doe",
            "prenom": "John",
            "email": "john.doe@example.com"
        },
        "error": null
    }
    """
    try:
        cv_generator = CVGenerator(current_user.id)
        user_data = cv_generator.get_user_data()

        if "error" in user_data:
            return (
                jsonify(
                    {"success": False, "message": user_data["error"], "data": user_data}
                ),
                404,
            )

        return jsonify({"success": True, "data": user_data})

    except Exception as e:
        logger.error(f"Erreur lors de la récupération des données: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la récupération des données: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/optimize", methods=["POST"])
@login_required
def optimize_cv_content():
    """
    Optimise une section du CV avec l'IA.

    Cette fonction est appelée lorsque l'utilisateur souhaite optimiser une section
    de son CV à l'aide d'une IA. Elle prend en entrée une requête POST contenant
    les informations nécessaires à l'optimisation, et renvoie une réponse JSON
    contenant les informations suivantes :
    - success (bool): Indique si la requête a été effectuée avec succès
    - original (str): Contenu original de la section avant optimisation
    - optimized (str): Contenu optimisé de la section par l'IA

    En cas d'erreur, la réponse JSON contiendra les informations suivantes :
    - success (bool): Indique si la requête a été effectuée avec succès
    - message (str): Message d'erreur

    Exemple de réponse JSON en cas de succès :
    {
        "success": true,
        "original": "Contenu original",
        "optimized": "Contenu optimisé"
    }

    Exemple de réponse JSON en cas d'erreur :
    {
        "success": false,
        "message": "Message d'erreur"
    }
    """
    try:
        data = request.get_json()
        section = data.get("section")
        content = data.get("content")

        if not section or not content:
            return (
                jsonify({"success": False, "message": "Section et contenu requis"}),
                400,
            )

        cv_generator = CVGenerator(current_user.id)

        if not cv_generator.gemini:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Service d'optimisation IA non disponible",
                    }
                ),
                503,
            )

        # Générer une version optimisée
        optimized = cv_generator._generate_content_with_gemini(
            section, {"content": content}
        )

        return jsonify({"success": True, "original": content, "optimized": optimized})

    except Exception as e:
        logger.error(f"Erreur lors de l'optimisation: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de l'optimisation: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/export/history", methods=["GET"])
@login_required
def get_export_history():
    """
    Récupère l'historique des CV exportés.

    Cette fonction est appelée lorsque l'utilisateur souhaite obtenir la liste des CV
    qu'il a générés. Elle renvoie une réponse JSON contenant les informations suivantes :

    - success (bool): Indique si la requête a été effectuée avec succès
    - user_exports (list): Liste des CV exportés par l'utilisateur, avec les
      informations suivantes pour chaque élément :
      - filename (str): Nom du fichier
      - format (str): Format du fichier (PDF, DOCX, etc.)
      - size (int): Taille du fichier en octets
      - created_at (str): Date et heure de création du fichier au format ISO 8601
      - download_url (str): Lien vers le fichier pour le téléchargement

    Exemple de réponse JSON en cas de succès :
    {
        "success": true,
        "user_exports": [
            {
                "filename": "CV_John_Doe_2022-01-01_10-00-00.pdf",
                "format": "PDF",
                "size": 1024,
                "created_at": "2022-01-01T10:00:00",
                "download_url": "/api/career-craft/download?path=CV_John_Doe_2022-01-01_10-00-00.pdf"
            },
            {
                "filename": "CV_John_Doe_2022-02-01_10-00-00.docx",
                "format": "DOCX",
                "size": 1536,
                "created_at": "2022-02-01T10:00:00",
                "download_url": "/api/career-craft/download?path=CV_John_Doe_2022-02-01_10-00-00.docx"
            }
        ]
    }

    Exemple de réponse JSON en cas d'erreur :
    {
        "success": false,
        "message": "Message d'erreur"
    }
    """
    try:
        user_exports = []

        # Parcourir les fichiers exportés
        for file_path in OUTPUT_FOLDER.glob(
            f"CV_{current_user.prenom}_{current_user.nom}_*.pdf"
        ):
            stat = file_path.stat()
            user_exports.append(
                {
                    "filename": file_path.name,
                    "format": file_path.suffix[1:].upper(),
                    "size": stat.st_size,
                    "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "download_url": f"/api/career-craft/download?path={file_path.relative_to(current_app.root_path)}",
                }
            )

        # Trier par date de création (plus récent en premier)
        user_exports.sort(key=lambda x: x["created_at"], reverse=True)

        return jsonify(
            {"success": True, "exports": user_exports, "total": len(user_exports)}
        )

    except Exception as e:
        logger.error(
            f"Erreur lors de la récupération de l'historique: {e}", exc_info=True
        )
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la récupération de l'historique: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/delete/<filename>", methods=["DELETE"])
@login_required
def delete_export(filename: str):
    """
    Supprime un fichier d'export du CV.

    Cette fonction est appelée lorsque l'utilisateur souhaite supprimer un fichier
    d'export de son CV. Elle prend en entrée le nom du fichier à supprimer.

    Si le fichier n'appartient pas à l'utilisateur, la fonction renvoie une réponse
    JSON avec un message d'erreur et un code d'état 403.

    Si le fichier n'est pas trouvé, la fonction renvoie une réponse JSON avec
    un message d'erreur et un code d'état 404.

    Si la suppression réussit, la fonction renvoie une réponse JSON avec un message
    de succès et un code d'état 200.

    Args:
        filename (str): Nom du fichier à supprimer.

    Returns:
        Une réponse JSON contenant les informations suivantes :
        - success (bool): Indique si la requête a été effectuée avec succès
        - message (str): Message de réussite ou d'erreur
        - code (int): Code d'état de la réponse

    Exemple de réponse JSON pour une suppression réussie :
    {
        "success": true,
        "message": "Fichier supprimé avec succès",
        "code": 200
    }
    """
    try:
        # Vérifier que le fichier appartient à l'utilisateur
        if not filename.startswith(f"CV_{current_user.prenom}_{current_user.nom}_"):
            return jsonify({"success": False, "message": "Accès non autorisé"}), 403

        file_path = OUTPUT_FOLDER / filename

        if not file_path.exists():
            return jsonify({"success": False, "message": "Fichier non trouvé"}), 404

        # Supprimer le fichier
        file_path.unlink()

        logger.info(f"Fichier supprimé: {filename}")

        return jsonify({"success": True, "message": "CV supprimé avec succès"})

    except Exception as e:
        logger.error(f"Erreur lors de la suppression: {e}", exc_info=True)
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la suppression: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/suggestions", methods=["POST"])
@login_required
def get_cv_suggestions():
    """
    Obtient des suggestions d'amélioration pour le CV.

    Cette fonction est appelée lorsque l'utilisateur souhaite obtenir des suggestions
    d'amélioration pour son CV. Elle utilise le modèle de génération de CV Gemini pour
    analyser le CV actuel de l'utilisateur et fournir des suggestions concrètes et
    actionnables.

    Si le service de suggestions n'est pas disponible, la fonction renvoie une réponse
    JSON avec un message d'erreur et un code d'état 503.

    Si la récupération des données du CV échoue, la fonction renvoie une réponse JSON avec
    un message d'erreur et un code d'état 404.

    Si la requête est valide et les données du CV sont correctement récupérées, la fonction
    renvoie une réponse JSON contenant les informations suivantes :
    - success (bool): Indique si la requête a été effectuée avec succès
    - suggestions (dict): Dictionnaire contenant les suggestions d'amélioration par catégorie
        - contenu (list): Liste de suggestions de contenu
        - structure (list): Liste de suggestions de structure et d'organisation
        - formulation (list): Liste de suggestions de formulation et de phraséologie
        - compétences (list): Liste de suggestions de compétences à ajouter ou mettre en valeur
        - impact (list): Liste de suggestions pour rendre le CV plus percutant

    Exemple de réponse JSON pour une requête réussie :
    {
        "success": true,
        "suggestions": {
            "contenu": [
                "Ajouter une expérience dans un domaine technique spécifique",
                "Mettre en évidence une compétence en plus du profil",
                "Ajouter une référence clientèle"
            ],
            "structure": [
                "Organiser les expériences professionnelles par ordre chronologique",
                "Mettre en évidence les compétences spécifiques dans chaque expérience",
                "Mettre en valeur les compétences dans la section compétences"
            ],
            "formulation": [
                "Utiliser des vocabulaire plus précis pour les compétences",
                "Éviter les mauvaises phrasés et les mauvaises constructions",
                "Utiliser une ponctuation appropriée"
            ],
            "compétences": [
                "Ajouter une compétence en programmation",
                "Mettre en valeur une compétence spécifique",
                "Ajouter une compétence en gestion de projet"
            ],
            "impact": [
                "Utiliser des titres attrayants pour les sections",
                "Mettre en valeur les expériences et les compétences",
                "Ajouter une section de motivations"
            ]
        }
    }
    """
    try:
        cv_generator = CVGenerator(current_user.id)

        if not cv_generator.gemini:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Service de suggestions non disponible",
                    }
                ),
                503,
            )

        user_data = cv_generator.get_user_data()

        if "error" in user_data:
            return jsonify({"success": False, "message": user_data["error"]}), 404

        # Créer un prompt pour obtenir des suggestions
        prompt = f"""
Analyse ce CV et fournis des suggestions d'amélioration concrètes et actionnables.

Données du CV:
{json.dumps(user_data, indent=2, ensure_ascii=False, default=str)}

Fournis tes suggestions dans les catégories suivantes:
1. Contenu: Ce qui manque ou pourrait être amélioré
2. Structure: Organisation et présentation
3. Formulation: Phrases et vocabulaire
4. Compétences: Compétences à ajouter ou mettre en valeur
5. Impact: Comment rendre le CV plus percutant

Format ta réponse en JSON avec cette structure:
{{
    "score_global": <note sur 100>,
    "points_forts": ["point 1", "point 2", ...],
    "points_amelioration": ["point 1", "point 2", ...],
    "suggestions": {{
        "contenu": ["suggestion 1", ...],
        "structure": ["suggestion 1", ...],
        "formulation": ["suggestion 1", ...],
        "competences": ["suggestion 1", ...],
        "impact": ["suggestion 1", ...]
    }}
}}
"""

        response = cv_generator.gemini.generate_response(prompt, temperature=0.5)

        if not response or not response.get("success"):
            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Impossible de générer des suggestions",
                    }
                ),
                500,
            )

        # Parser la réponse JSON
        suggestions_text = response.get("response", "{}")

        # Nettoyer le texte (supprimer les balises markdown si présentes)
        suggestions_text = re.sub(r"```json\n?", "", suggestions_text)
        suggestions_text = re.sub(r"```\n?", "", suggestions_text)

        try:
            suggestions = json.loads(suggestions_text)
        except json.JSONDecodeError:
            logger.error(f"Impossible de parser les suggestions: {suggestions_text}")
            suggestions = {
                "score_global": 70,
                "points_forts": ["CV bien structuré"],
                "points_amelioration": ["Ajouter plus de détails"],
                "suggestions": {
                    "contenu": ["Développez vos expériences"],
                    "structure": ["Organisez par ordre chronologique"],
                    "formulation": ["Utilisez des verbes d'action"],
                    "competences": ["Ajoutez des compétences techniques"],
                    "impact": ["Quantifiez vos réalisations"],
                },
            }

        return jsonify({"success": True, "suggestions": suggestions})

    except Exception as e:
        logger.error(
            f"Erreur lors de la génération des suggestions: {e}", exc_info=True
        )
        return (
            jsonify(
                {
                    "success": False,
                    "message": f"Erreur lors de la génération des suggestions: {str(e)}",
                }
            ),
            500,
        )


@career_craft_bp.route("/health", methods=["GET"])
def health_check():
    """
    Endpoint de vérification de santé du service

    Cette fonction permet de vérifier si le service est sain et opérationnel.
    Elle renvoie une réponse JSON contenant les informations suivantes :
    - success (bool) : Indique si la vérification a été effectuée avec succès

    Si la vérification réussit, la réponse JSON contient également les informations suivantes :
    - folders_ok (bool) : Indique si les dossiers nécessaires existent
    - template_count (int) : Nombre de modèles de CV disponibles
    - gemini_available (bool) : Indique si Gemini est disponible
    - gemini_api_key_set (bool) : Indique si la clé API de Gemini est définie

    Si la vérification échoue, la réponse JSON contient également les informations suivantes :
    - message (str) : Message d'erreur détaillé
    """
    try:
        # Vérifier que les dossiers nécessaires existent
        folders_ok = all(
            [UPLOAD_FOLDER.exists(), TEMPLATES_FOLDER.exists(), OUTPUT_FOLDER.exists()]
        )

        # Compter les templates disponibles
        template_count = len(list(TEMPLATES_FOLDER.glob("*.json")))

        # Vérifier la disponibilité de Gemini
        gemini_api_key = os.getenv("GEMINI_API_KEY")
        gemini_available = bool(gemini_api_key)

        return jsonify(
            {
                "success": True,
                "status": "healthy" if folders_ok else "degraded",
                "checks": {
                    "folders": folders_ok,
                    "templates_available": template_count,
                    "gemini_available": gemini_available,
                },
                "version": "2.0.0",
                "timestamp": datetime.now().isoformat(),
            }
        )

    except Exception as e:
        logger.error(f"Erreur lors du health check: {e}", exc_info=True)
        return jsonify({"success": False, "status": "unhealthy", "error": str(e)}), 500


# ============================================================================
# UTILITAIRES
# ============================================================================


def cleanup_old_exports(days: int = 7):
    """Nettoie les exports de plus de X jours"""
    try:
        cutoff_time = datetime.now().timestamp() - (days * 24 * 3600)
        deleted_count = 0

        for file_path in OUTPUT_FOLDER.glob("CV_*.pdf"):
            if file_path.stat().st_ctime < cutoff_time:
                file_path.unlink()
                deleted_count += 1
                logger.info(f"Fichier supprimé (ancien): {file_path.name}")

        for file_path in OUTPUT_FOLDER.glob("CV_*.docx"):
            if file_path.stat().st_ctime < cutoff_time:
                file_path.unlink()
                deleted_count += 1
                logger.info(f"Fichier supprimé (ancien): {file_path.name}")

        logger.info(f"Nettoyage terminé: {deleted_count} fichiers supprimés")
        return deleted_count

    except Exception as e:
        logger.error(f"Erreur lors du nettoyage: {e}", exc_info=True)
        return 0


# Enregistrer la tâche de nettoyage (peut être appelée par un scheduler)
@career_craft_bp.route("/admin/cleanup", methods=["POST"])
@login_required
def trigger_cleanup():
    """
    Déclenche le nettoyage des anciens exports.

    Cette fonction est réservée aux administrateurs et permet de nettoyer les exports
    anciens de plus de X jours.

    Paramètres:
        - Aucun

    Retour:
        - Une réponse JSON contenant:
            - success (bool): Indique si la requête a été effectuée avec succès
            - message (str): Message de réussite ou d'erreur
            - deleted_count (int): Nombre de fichiers supprimés

    """
    try:
        # Vérifier que l'utilisateur est admin
        if not hasattr(current_user, "role") or current_user.role != "admin":
            return jsonify({"success": False, "message": "Accès non autorisé"}), 403

        data = request.get_json() or {}
        days = data.get("days", 7)

        deleted_count = cleanup_old_exports(days)

        return jsonify(
            {
                "success": True,
                "message": f"{deleted_count} fichiers supprimés",
                "deleted_count": deleted_count,
            }
        )

    except Exception as e:
        logger.error(f"Erreur lors du nettoyage manuel: {e}", exc_info=True)
        return (
            jsonify(
                {"success": False, "message": f"Erreur lors du nettoyage: {str(e)}"}
            ),
            500,
        )


# ============================================================================
# ERROR HANDLERS
# ============================================================================


@career_craft_bp.errorhandler(404)
def not_found(error):
    """Gestionnaire d'erreur 404"""
    return (
        jsonify(
            {"success": False, "error": "Ressource non trouvée", "message": str(error)}
        ),
        404,
    )


@career_craft_bp.errorhandler(500)
def internal_error(error):
    """Gestionnaire d'erreur 500"""
    logger.error(f"Erreur interne: {error}", exc_info=True)
    return (
        jsonify(
            {
                "success": False,
                "error": "Erreur interne du serveur",
                "message": "Une erreur s'est produite. Veuillez réessayer.",
            }
        ),
        500,
    )


# ============================================================================
# INITIALISATION
# ============================================================================


def init_career_craft(app):
    """Initialise le module CareerCraft avec l'application Flask"""
    app.register_blueprint(career_craft_bp)
    logger.info("Module CareerCraft initialisé avec succès")

    # Créer les dossiers nécessaires
    for folder in [UPLOAD_FOLDER, TEMPLATES_FOLDER, OUTPUT_FOLDER]:
        folder.mkdir(parents=True, exist_ok=True)

    logger.info("Dossiers créés: Upload, Templates, Output")


if __name__ == "__main__":
    # Tests unitaires basiques
    print("CareerCraft AI - Module de génération de CV")
    print("=" * 50)
    print(f"Dossier uploads: {UPLOAD_FOLDER}")
    print(f"Dossier templates: {TEMPLATES_FOLDER}")
    print(f"Dossier exports: {OUTPUT_FOLDER}")
    print(f"Gemini API configurée: {bool(os.getenv('GEMINI_API_KEY'))}")
    print("=" * 50)
